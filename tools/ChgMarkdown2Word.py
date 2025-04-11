import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from markdown import markdown
from bs4 import BeautifulSoup
import re

class MarkdownToWordConverter:
    def __init__(self, config_file):
        self.load_config(config_file)
        self.doc = Document()
        self._define_styles()
    
    def load_config(self, config_file):
        """加载格式配置文件（JSON格式）"""
        with open(config_file, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        
        # 设置默认字体
        font_config = self.config.get('font', {})
        self.default_font = font_config.get('name', '宋体')
        self.default_size = Pt(font_config.get('size', 12))
        self.default_color = self._hex_to_rgb(font_config.get('color', '#000000'))

    def _define_styles(self):
        """预定义Word样式"""
        styles = self.doc.styles
        
        # 标题样式
        for level in ['h1', 'h2', 'h3']:
            if level in self.config.get('headings', {}):
                config = self.config['headings'][level]
                style = styles.add_style(f"Heading {level}", 1)
                font = style.font
                font.name = config.get('font', self.default_font)
                font.size = Pt(config.get('size', 16 if level == 'h1' else 14))
                font.color.rgb = self._hex_to_rgb(config.get('color', '#000000'))
                font.bold = config.get('bold', True)
                
                # 段落格式
                paragraph_format = style.paragraph_format
                paragraph_format.space_before = Pt(config.get('space_before', 12))
                paragraph_format.space_after = Pt(config.get('space_after', 6))
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 'h1' else WD_PARAGRAPH_ALIGNMENT.LEFT

        # 正文样式
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = self.default_font
        font.size = self.default_size
        font.color.rgb = self.default_color

    def _hex_to_rgb(self, hex_color):
        """将十六进制颜色转换为RGB"""
        hex_color = hex_color.lstrip('#')
        return RGBColor(*tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4)))

    def _process_paragraph(self, element):
        """处理段落格式"""
        p = self.doc.add_paragraph()
        
        # 应用段落格式
        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.3)  # 首行缩进2字符
        p_format.line_spacing = 1.5  # 1.5倍行距
        p_format.space_after = Pt(6)

        # 处理内联样式
        self._parse_inline_styles(p, element)

    def _parse_inline_styles(self, paragraph, element):
        """解析内联样式（粗体、斜体等）"""
        for content in element.contents:
            if content.name == 'strong':
                self._add_text_run(paragraph, content.text, bold=True)
            elif content.name == 'em':
                self._add_text_run(paragraph, content.text, italic=True)
            elif content.name == 'a':
                self._add_hyperlink(paragraph, content['href'], content.text)
            else:
                self._add_text_run(paragraph, str(content))

    def _add_text_run(self, paragraph, text, **kwargs):
        """添加带格式的文字"""
        run = paragraph.add_run(text)
        run.font.name = self.default_font
        run.font.size = self.default_size
        
        # 应用格式参数
        for key, value in kwargs.items():
            setattr(run.font, key, value)

    def _add_hyperlink(self, paragraph, url, text):
        """添加超链接"""
        part = paragraph.part
        r_id = part.relate_to(
            url, 
            'hyperlink',
            is_external=True
        )

        hyperlink = paragraph.add_run()
        hyperlink._r.append(self._create_hyperlink_element(r_id, text))

    def _create_hyperlink_element(self, r_id, text):
        """创建超链接XML元素"""
        return BeautifulSoup(f'''
            <w:hyperlink r:id="{r_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:r>
                    <w:rPr>
                        <w:rStyle w:val="Hyperlink"/>
                        <w:color w:val="0000FF"/>
                        <w:u w:val="single"/>
                    </w:rPr>
                    <w:t>{text}</w:t>
                </w:r>
            </w:hyperlink>
        ''', 'xml').find('w:hyperlink')

    def convert(self, markdown_text, output_file):
        """执行转换"""
        # 转换Markdown为HTML
        html = markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')
        
        # 遍历HTML元素
        for element in soup:
            if element.name in ['h1', 'h2', 'h3']:
                # 处理标题
                heading = self.doc.add_heading(level=int(element.name))
                heading.add_run(element.text).bold = True
                heading.style = f"Heading {element.name}"
            elif element.name == 'p':
                # 处理段落
                self._process_paragraph(element)
            elif element.name in ['ul', 'ol']:
                # 处理列表
                for li in element.find_all('li'):
                    self.doc.add_paragraph(li.text, style='ListBullet' if element.name == 'ul' else 'ListNumber')
            elif element.name == 'hr':
                # 添加分页符
                self.doc.add_page_break()
        
        # 保存文档
        self.doc.save(output_file)

# 示例配置文件格式（config.json）
"""
{
    "font": {
        "name": "宋体",
        "size": 12,
        "color": "#333333"
    },
    "headings": {
        "h1": {
            "size": 22,
            "color": "#2A5B87",
            "bold": true,
            "space_before": 24,
            "space_after": 12
        },
        "h2": {
            "size": 18,
            "color": "#3F7FBF",
            "space_before": 18,
            "space_after": 9
        }
    }
}
"""

# 使用示例
if __name__ == "__main__":
    # 示例Markdown文本
    sample_md = """
# 人工智能报告

## 技术发展

==&zwnj;**自然语言处理**&zwnj;==领域近年来取得重大突破，主要技术包括：

- 预训练语言模型（如BERT）
- 注意力机制
- 生成对抗网络

[查看最新论文](https://arxiv.org/)

---

### 未来趋势
预计到2025年，AI将实现更*自然*的对话能力。
"""

    # 执行转换
    converter = MarkdownToWordConverter('config.json')
    converter.convert(sample_md, 'output.docx')
